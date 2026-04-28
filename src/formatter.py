from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path


SECTION_HEADERS = [
    "=== FILE NOTES ===",
    "=== TIME RECORDING ===",
    "=== CORRESPONDENCE ===",
    "=== ACTION ITEMS ===",
]


def parse_sections(text):
    sections = {}
    current_header = "TRANSCRIPT"
    current_lines = []

    for line in text.split("\n"):
        stripped = line.strip()
        # Normalize "### === HEADER ===" to "=== HEADER ==="
        normalized = stripped.lstrip("# ")
        if normalized in SECTION_HEADERS:
            if current_lines:
                sections[current_header] = "\n".join(current_lines).strip()
            current_header = normalized.strip("= ")
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections[current_header] = "\n".join(current_lines).strip()

    return sections


def _set_run_font(run, name="Times New Roman", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def build_docx(raw_text, corrected_text, formatted_text, source_file):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15

    p = doc.add_paragraph()
    run = p.add_run("Legal Dictation Transcription")
    _set_run_font(run, bold=True, size=14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f"Date: {date.today().isoformat()}")
    _set_run_font(run, size=10)
    run = p.add_run(f"\nSource: {Path(source_file).name}")
    _set_run_font(run, size=10)

    doc.add_paragraph().add_run("").add_break()

    sections = parse_sections(formatted_text or corrected_text or raw_text)

    for header, body in sections.items():
        p = doc.add_paragraph()
        run = p.add_run(header)
        _set_run_font(run, bold=True, size=12)

        if header == "ACTION ITEMS":
            for line in body.split("\n"):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                checkbox = "[ ] " if not line.startswith("[") else ""
                run = p.add_run(f"{checkbox}{line}")
                _set_run_font(run, size=12)
        else:
            paragraphs = body.split("\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(para)
                _set_run_font(run, size=12)

    return doc


def save_docx(doc, output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
